from bs4 import BeautifulSoup
import requests
from datetime import date, timedelta, datetime
from googletrans import Translator
from docx import Document
from docx.shared import Inches
import os
from io import BytesIO
from PIL import Image


#Get last 7 day and insert to array
today = date.today()
lst_7day = date.today() - timedelta(days=7)
delta = today - lst_7day   # returns timedelta
lst7days_array = []
for i in range(delta.days + 1):
    day = lst_7day + timedelta(days=i)
    lst7days_array.append(day.strftime("%B %d, %Y"))
#print(lst7days_array)



def getRef():
	time_now= datetime.now()
	# Time with format 2022-03-23T16:01:14
	timeDate = time_now.strftime("%Y-%m-%dT%H:%M:%S")
	
	# Get first page and next page
	html1 = requests.get('https://thehackernews.com/search?updated-max='+timeDate+'&max-results=20')
	soup = BeautifulSoup(html1.text, "html.parser")
	a = soup.find(class_ ='blog-pager-older-link-mobile')
	html2 = requests.get(a['href'])
	html = html1.text+html2.text
	soup = BeautifulSoup(html, "html.parser")

	#Get ref of aritcle last 7 days 
	ref_arr = []
	for b in soup.find_all('div' ,class_='body-post clear'):
		publish_date = b.find(class_='item-label').text.replace('','Date: ',).replace('','Author')
		for k in lst7days_array:
			if k in publish_date:
				#title = b.find(class_='home-title').text
				link_ref = b.find(class_='story-link')
				#print(title + publish_date)
				ref_arr.append(link_ref['href'])
	return ref_arr


def getArticle():
	translator = Translator()
	document = Document()
	document.add_heading('Tin tức an ninh mạng từ '+lst_7day.strftime("%d-%m-%Y")+' tới '+today.strftime("%d-%m-%Y"), 0)
	for i in getRef():
		html = requests.get(i)
		soup = BeautifulSoup(html.text, "html.parser")	
		for k in soup.find_all('div',class_='main-box clear'):
			title = k.find('h1',class_='story-title').text
			publish_date_author = k.find(class_='postmeta').text.replace('','📅Date: ',).replace('',' 👤Author: ')
			image_link = k.find('img')['src']
			detail = k.find(class_='articlebody clear cf').text.replace('','',)
			titleTranlate = translator.translate(text=title,src='en',dest='vi').text
			detailTranlate = translator.translate(text=detail,src='en',dest='vi').text
			#Write to TXT file
			with open('THN.txt',"a+", encoding='UTF-8') as f:		
				f.write('Title: 🐄'+title+'\n')
				f.write(publish_date_author+'\n')
				f.write('🌎Image:'+image_link+'\n')
				f.write(detail)
				f.write('Reference Link:'+i)
				f.write('\n'+'======================================================================================'+'\n')
				f.write('🕮Translate to Vietnamese: '+'\n')
				f.write('Title: 🐄'+titleTranlate+'\n'+'\n')
				f.write(detailTranlate)
				f.write('\n'+'++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++'+'\n\n')
			print('Success write to txt file')
			#Write to M$Word
			response = requests.get(image_link)
			image = Image.open(BytesIO(response.content))
			try:
				image.save('image.jpg')		
			except:
				with open("image.jpg", 'wb') as f:
					f.write(response.content)
			document.add_heading(title, level=1)
			document.add_paragraph(publish_date_author)
			p = document.add_paragraph('Link tham khảo: ')
			p.add_run(i).italic = True
			document.add_picture('image.jpg',width=Inches(6.0))
			table = document.add_table(rows=2, cols=2)
			table.style = 'Table Grid'
			hdr_cells = table.rows[0].cells
			hdr_cells[0].text = 'Title: 🐄'+title
			hdr_cells[1].text = 'Title: 🐄'+titleTranlate
			row_cells = table.rows[1].cells
			row_cells[0].text = detail
			row_cells[1].text = detailTranlate
			os.remove('image.jpg')
			print('Success write to Word')
	document.save('THN.docx')


getArticle()

